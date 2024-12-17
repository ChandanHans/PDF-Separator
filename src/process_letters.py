from datetime import datetime
import os
import tempfile
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter

from .process_labels import create_labels
from .drive_upload import download_image, get_table_data
from .utils import execute_with_retry, extract_number, normalize_rows

from .constants import *

def create_pdf_from_template(replacements, image_path, output_folder, pdf_filename):
    # Load the Word document template
    doc = Document(LETTER_TEMPLATE)

    # Function to modify text in a run
    def modify_run(run):
        for find_text, replace_text in replacements.items():
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)

    # Replace text in the header
    header = doc.sections[0].header
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            modify_run(run)

    # Replace text in the body
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            modify_run(run)

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        modify_run(run)

    # Ensure the image spans a full page
    section = doc.sections[0]
    page_width = section.page_width - section.left_margin
    page_height = section.page_height - section.top_margin

    # Add the image to the document and adjust size to full page
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=page_width, height=page_height)

    # Create a temporary file to store the intermediate Word document
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_doc:
        temp_doc_path = temp_doc.name
        doc.save(temp_doc_path)

    # Define PDF output path
    pdf_output_path = os.path.join(output_folder, pdf_filename)
    convert(temp_doc_path, pdf_output_path)

    # Remove the temporary Word document
    os.remove(temp_doc_path)

def combine_pdfs(input_folder, output_folder, output_filename):
    """
    Combines all PDF files in a folder into a single PDF, saves it to another folder, and deletes the original PDFs.

    :param input_folder: Path to the folder containing PDF files to combine.
    :param output_folder: Path to the folder where the combined PDF will be saved.
    :param output_filename: Name of the combined PDF file (e.g., 'combined.pdf').
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)  # Create the output folder if it doesn't exist
    
    pdf_writer = PdfWriter()
    pdf_files = []  # To track PDF files for deletion

    files_in_folder = sorted(os.listdir(input_folder), key = extract_number)
    if not files_in_folder:
        print("There is no files in Temp-Letter to combine")
        return
    
    # Iterate through all files in the input folder
    for file_name in files_in_folder:
        file_path = os.path.join(input_folder, file_name)
        
        # Check if the file is a PDF
        if file_name.lower().endswith('.pdf'):
            try:
                pdf_reader = PdfReader(file_path)
                pdf_files.append(file_path)  # Add the file path to the list for deletion
                # Add all pages from the current PDF to the writer
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
            except Exception as e:
                print(f"Error reading {file_name}: {e}")
    
    # Save the combined PDF to the output folder
    output_path = os.path.join(output_folder, output_filename)
    with open(output_path, 'wb') as output_file:
        pdf_writer.write(output_file)

    print(f"Combined PDF saved at: {output_path}")

    for pdf_file in pdf_files:
        try:
            os.remove(pdf_file)
        except Exception as e:
            pass

def create_combine_letters(sheets_service, drive_service):
    today = datetime.now().date()
    new_date_text = today.strftime("%d-%b-%Y")
    
    all_values = get_table_data(sheets_service, ANNUAIRE_HERITIERS_SHEET_ID, "Héritier Annuaire!A:N")

    # Normalize all rows to ensure they have 8 elements
    normalized_values : list[str] = normalize_rows(all_values[1:], 14)
    labels = []
    for index, row in enumerate(normalized_values, start=2):
        if not row[10] or row[10] == "Not contacted" and row[13] == "Vérifié":
            print(index)
            name = row[0]
            image_link = row[9]
            replacements = {'(NAME)': name}
            latter_file_name = f"Letter - {index}.pdf"
            image = download_image(drive_service, image_link)
            create_pdf_from_template(replacements, image, TEMP_LETTER_FOLDER, latter_file_name)
            
            request = sheets_service.spreadsheets().values().update(
                spreadsheetId=ANNUAIRE_HERITIERS_SHEET_ID,
                range=f"Héritier Annuaire!K{index}:L{index}",
                valueInputOption="USER_ENTERED",  # Allows typing-like behavior
                body={"values": [["Contacted / pending answer", new_date_text]]}
            )
            execute_with_retry(request)
            heir_full_name : str = row[3]
            city = row[5].split("(")[0]
            heir_name = "".join((heir_full_name.split()[0],heir_full_name.split()[-1]))
            labels.append((heir_name,row[4],row[6],city))
            
    current_time = datetime.now()
    formatted_time = current_time.strftime("%d-%m-%Y-%H-%M")
    combine_pdfs(TEMP_LETTER_FOLDER, LETTER_FOLDER, f"Letters - {formatted_time}.pdf")
    create_labels(labels, LETTER_FOLDER, f"Labels - {formatted_time}.pdf")