import os
import tempfile
from docx import Document
from docx2pdf import convert

from .constants import *

def create_pdf_from_template(replacements, image_path, pdf_filename):
    # Load the Word document template
    doc = Document(LETTER_TEMPLATE)

    # Function to modify text in a run
    def modify_run(run):
        for find_text, replace_text in replacements.items():
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)

    # Replace text in the header
    header = doc.sections[0].header
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            modify_run(run)

    # Replace text in the body
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            modify_run(run)

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        modify_run(run)

    # Ensure the image spans a full page
    section = doc.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin
    page_height = section.page_height - section.top_margin - section.bottom_margin

    # Add the image to the document and adjust size to full page
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=page_width, height=page_height)

    # Create a temporary file to store the intermediate Word document
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_doc:
        temp_doc_path = temp_doc.name
        doc.save(temp_doc_path)

    # Define PDF output path
    pdf_output_path = os.path.join(LETTER_FOLDER, pdf_filename)
    convert(temp_doc_path, pdf_output_path)

    # Remove the temporary Word document
    os.remove(temp_doc_path)